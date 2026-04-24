import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from tkinterdnd2 import DND_FILES, TkinterDnD
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import spacy
import re
import os
import sys
import subprocess
import threading
import difflib
from datetime import datetime, timedelta
import copy
import traceback
import hashlib
import uuid
import json
import base64

# 【终极打包防漏件修复】：多重兼容加载机制
try:
    import zh_core_web_sm
    # 优先尝试标准库加载（最稳定）
    try:
        nlp = spacy.load("zh_core_web_sm")
    except:
        # 如果打包后标准加载失败，尝试直接调用包内置的加载方法
        nlp = zh_core_web_sm.load()
except Exception as e:
    err_msg = traceback.format_exc()
    messagebox.showerror(
        "打包环境缺陷 (缺少底层字典)", 
        f"叙叙，AI 模型底层的分词字典 (spacy_pkuseg) 被打包工具漏掉了！\n\n请删除 build、dist 文件夹，然后务必使用下面这行包含 '--collect-all spacy_pkuseg' 的最新口令重新打包：\n\npyinstaller --noconsole --onefile --collect-all tkinterdnd2 --collect-all spacy --collect-all zh_core_web_sm --collect-all spacy_pkuseg --copy-metadata spacy --copy-metadata zh_core_web_sm --copy-metadata spacy_pkuseg --icon=NONE AutoRedactor.py\n\n底层报错信息截取:\n{err_msg[-300:]}"
    )
    sys.exit()

# ==================== 正则表达式定义 (规则匹配) ====================

# 0. 锚点前缀规则 (用于动态记忆提取，保留贪婪匹配以获取简称)
PREFIX_REGEX = re.compile(r'((?:甲\s*方|乙\s*方|丙\s*方|丁\s*方|戊\s*方|供\s*方|需\s*方|出\s*卖\s*人|买\s*受\s*人|委\s*托\s*方|受\s*托\s*方|委\s*托\s*人|受\s*托\s*人|加\s*工\s*方|承\s*揽\s*方|居\s*间\s*方|服\s*务\s*方|顾\s*问\s*方|授\s*权\s*方|被\s*授\s*权\s*方|出\s*租\s*方|承\s*租\s*方|法\s*定\s*地\s*址|注\s*册\s*地\s*址|办\s*公\s*地\s*址|实\s*际\s*地\s*址|地\s*址|住\s*址|住\s*所(?:\s*地)?|居\s*住\s*地|签\s*订\s*地\s*点|签\s*署\s*地\s*点|交\s*货\s*地\s*点|联\s*系\s*人|联\s*系\s*地\s*址|联\s*系\s*方\s*式|联\s*系\s*电\s*话|收\s*款\s*人|收\s*款\s*[账帐]\s*号|收\s*款\s*银\s*行|法\s*定\s*代\s*表\s*人|法\s*人\s*或\s*授\s*权\s*代\s*表(?:\s*人)?|法\s*人\s*代\s*表|法\s*人|负\s*责\s*人|授\s*权\s*代\s*表(?:\s*人)?|委\s*托\s*代\s*理\s*人|电\s*话|传\s*真|邮\s*箱|电\s*子\s*信\s*箱|电\s*子\s*邮\s*件|电\s*邮|邮\s*政\s*编\s*码|邮\s*编|微\s*信(?:\s*号)?|Q\s*Q(?:\s*号)?|支\s*付\s*宝(?:\s*[账帐]\s*号)?|昵\s*称|姓\s*名|身\s*份\s*证(?:\s*件)?(?:\s*号\s*码|\s*号)?|公\s*民\s*身\s*份\s*号\s*码|签\s*发\s*机\s*关|[账帐]\s*号|开\s*户\s*(?:银\s*行|行)|(?:银\s*行\s*)?卡\s*号|银\s*行\s*[账帐]\s*户|银\s*行\s*[账帐]\s*号|[账帐]\s*户\s*名\s*称|[账帐]\s*户\s*号\s*码|户\s*名|大\s*写(?:\s*金\s*额)?|统\s*一\s*社\s*会\s*信\s*用\s*代\s*码|组\s*织\s*机\s*构\s*代\s*码|纳\s*税\s*人\s*识\s*别\s*号|税\s*号)\s*(?:[（(][^)）]+[)）])?\s*[：:]\s*(?:[\[【(（]\s*)?)\s*([^\n]+)')

# 0.5.1 隐藏专用的地址类锚点
HIDE_ADDRESS_REGEX = re.compile(r'((?:法\s*定\s*地\s*址|注\s*册\s*地\s*址|办\s*公\s*地\s*址|实\s*际\s*地\s*址|地\s*址|住\s*址|住\s*所(?:\s*地)?|居\s*住\s*地|签\s*订\s*地\s*点|签\s*署\s*地\s*点|交\s*货\s*地\s*点|联\s*系\s*地\s*址)\s*(?:[（(][^)）]+[)）])?\s*[：:]\s*(?:[\[【(（]\s*)?)\s*([^\n，。；;,.;、]+)')

# 0.5.2 隐藏专用的短文本锚点
HIDE_SHORT_REGEX = re.compile(r'((?:甲\s*方|乙\s*方|丙\s*方|丁\s*方|戊\s*方|供\s*方|需\s*方|出\s*卖\s*人|买\s*受\s*人|委\s*托\s*方|受\s*托\s*方|委\s*托\s*人|受\s*托\s*人|加\s*工\s*方|承\s*揽\s*方|居\s*间\s*方|服\s*务\s*方|顾\s*问\s*方|授\s*权\s*方|被\s*授\s*权\s*方|出\s*租\s*方|承\s*租\s*方|联\s*系\s*人|联\s*系\s*方\s*式|联\s*系\s*电\s*话|收\s*款\s*人|收\s*款\s*[账帐]\s*号|收\s*款\s*银\s*行|法\s*定\s*代\s*表\s*人|法\s*人\s*或\s*授\s*权\s*代\s*表(?:\s*人)?|法\s*人\s*代\s*表|法\s*人|负\s*责\s*人|授\s*权\s*代\s*表(?:\s*人)?|委\s*托\s*代\s*理\s*人|电\s*话|传\s*真|邮\s*箱|电\s*子\s*信\s*箱|电\s*子\s*邮\s*件|电\s*邮|邮\s*政\s*编\s*码|邮\s*编|微\s*信(?:\s*号)?|Q\s*Q(?:\s*号)?|支\s*付\s*宝(?:\s*[账帐]\s*号)?|昵\s*称|姓\s*名|身\s*份\s*证(?:\s*件)?(?:\s*号\s*码|\s*号)?|公\s*民\s*身\s*份\s*号\s*码|签\s*发\s*机\s*关|[账帐]\s*号|开\s*户\s*(?:银\s*行|行)|(?:银\s*行\s*)?卡\s*号|银\s*行\s*[账帐]\s*户|银\s*行\s*[账帐]\s*号|[账帐]\s*户\s*名\s*称|[账帐]\s*户\s*号\s*码|户\s*名|大\s*写(?:\s*金\s*额)?|统\s*一\s*社\s*会\s*信\s*用\s*代\s*码|组\s*织\s*机\s*构\s*代\s*码|纳\s*税\s*人\s*识\s*别\s*号|税\s*号)\s*(?:[（(][^)）]+[)）])?\s*[：:]\s*(?:[\[【(（]\s*)?)\s*([^\n，。；;,、()（）\[\]【】\s]+(?:[ \t]+[a-zA-Z0-9_.-]+)*)')

# 叙叙专属优化：地毯式拦截正文里的公司和律所
COMPANY_LAWFIRM_REGEX = re.compile(r'([a-zA-Z0-9\u4e00-\u9fa5（）\(\)·&\-\s]{2,20}(?:有限公司|有限责任公司|股份有限公司|集团公司|合伙企业|律师事务所|分公司|支公司|子公司|总公司|代表处|办事处))')

# 1-7. 其他通用正则
PHONE_REGEX = re.compile(r'\b1[3-9]\d{9}\b')
TEL_REGEX = re.compile(r'\b0\d{2,3}-\d{7,8}\b')
ID_REGEX = re.compile(r'\b[1-9]\d{5}(18|19|20)\d{2}(0[1-9]|1[0-2])([0-2][1-9]|10|20|30|31)\d{3}[0-9Xx]\b')
EMAIL_REGEX = re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+')
# 叙叙专属优化：允许金额数字被各类括号包裹
MONEY_REGEX_RULES = re.compile(r'(?:人民币|RMB|[¥￥])\s*(?:[\[【(（]\s*\d+(?:,\d{3})*(?:\.\d+)?\s*[\]】)）]|\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:元|万元|千元|万)?|(?:[\[【(（]\s*\d+(?:,\d{3})*(?:\.\d+)?\s*[\]】)）]|\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:元|万元|千元|万)')
CHINESE_MONEY_REGEX = re.compile(r'(?:人民币)?\s*[零壹贰叁肆伍陆柒捌玖拾佰仟万亿一二三四五六七八九十百千万亿]+(?:圆|元|角|分|整|正)[零壹贰叁肆伍陆柒捌玖拾佰仟万亿一二三四五六七八九十百千万亿角分整正]*')
BANK_BRANCH_REGEX = re.compile(r'[\u4e00-\u9fa5]{2,20}(?:分行|支行)')
# 8. 匹配比例与费率
PERCENTAGE_REGEX = re.compile(r'(?:[百千万]分之[零壹贰叁肆伍陆柒捌玖拾佰仟万亿一二三四五六七八九十百千万亿点\d]+|\d+(?:\.\d+)?\s*[%‰])')


class LicenseManager:
    """管理试用期和激活码的核心类"""
    def __init__(self):
        # 混淆文件路径，保存在用户根目录下，防止轻易被发现删除
        self.license_file = os.path.join(os.path.expanduser('~'), '.fashi_tool_sys.dat')
        self.secret_salt = "FASHI_LAB_XUXU_2026_SECRET_SALT" # 绝对机密盐值，千万不要改！
        self.trial_days = 7
        self.is_activated = False
        self.days_left = 0
        self.load_or_create_license()

    def get_machine_code(self):
        # 获取底层机器码 (MAC地址) 并哈希化，生成类似 ABCD-EFGH-IJKL-MNOP 的格式
        mac = str(uuid.getnode())
        hash_str = hashlib.md5(mac.encode()).hexdigest().upper()
        return f"{hash_str[:4]}-{hash_str[4:8]}-{hash_str[8:12]}-{hash_str[12:16]}"

    def generate_expected_activation_code(self):
        # 根据机器码生成真正的激活码
        raw = f"{self.get_machine_code()}_{self.secret_salt}"
        hash_str = hashlib.sha256(raw.encode()).hexdigest().upper()
        return f"{hash_str[:4]}-{hash_str[4:8]}-{hash_str[8:12]}-{hash_str[12:16]}"

    def load_or_create_license(self):
        current_date_str = datetime.now().strftime("%Y-%m-%d")
        
        if not os.path.exists(self.license_file):
            # 第一次运行，创建隐藏的授权文件
            data = {"first_run": current_date_str, "activation_code": ""}
            self._save_file(data)
        
        # 读取授权文件
        data = self._read_file()
        if not data or "first_run" not in data:
            data = {"first_run": current_date_str, "activation_code": ""}
            self._save_file(data)

        first_run = datetime.strptime(data["first_run"], "%Y-%m-%d")
        saved_code = data.get("activation_code", "")

        # 检查是否已永久激活
        if saved_code == self.generate_expected_activation_code():
            self.is_activated = True
            self.days_left = 999
            return

        # 没激活，计算试用期剩余天数
        delta = datetime.now() - first_run
        self.days_left = max(0, self.trial_days - delta.days)
        # 防止系统时间被篡改倒流
        if delta.days < 0:
            self.days_left = 0

    def verify_and_activate(self, input_code):
        if input_code.strip().upper() == self.generate_expected_activation_code():
            data = self._read_file()
            data["activation_code"] = input_code.strip().upper()
            self._save_file(data)
            self.is_activated = True
            return True
        return False

    def _save_file(self, data_dict):
        # 简单将 JSON 转成 Base64 保存，防止用户用记事本直接打开改时间
        raw_str = json.dumps(data_dict)
        encoded_str = base64.b64encode(raw_str.encode('utf-8')).decode('utf-8')
        try:
            with open(self.license_file, 'w') as f:
                f.write(encoded_str)
        except:
            pass

    def _read_file(self):
        try:
            with open(self.license_file, 'r') as f:
                encoded_str = f.read()
            raw_str = base64.b64decode(encoded_str.encode('utf-8')).decode('utf-8')
            return json.loads(raw_str)
        except:
            return None


class RedactorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("智能 Word 文档脱敏工具 (跨平台最终版)")
        self.root.geometry("520x620") # 再次拉高一点空间留给底部的激活按钮
        self.root.configure(padx=20, pady=20)
        
        # 初始化授权管理器
        self.license_mgr = LicenseManager()
        
        self.file_path = ""
        self.white_dict_path = os.path.join(os.getcwd(), "脱敏白名单.txt")
        self.black_dict_path = os.path.join(os.getcwd(), "脱敏黑名单.txt")
        
        # UI 字体跨平台适配 (Mac 使用苹果专属 PingFang，Win 使用微软雅黑)
        self.sys_font = "PingFang SC" if sys.platform == "darwin" else "Microsoft YaHei"
        
        # 运行时动态记忆的主体名词
        self.dynamic_black_keywords = set()
        
        # 【核心护体金牌】：这些身份词拥有绝对特权，任何情况都不能被删掉或变成黑名单！
        # 叙叙专属修复：补充加入签字盖章动作的绝对免疫，防止被当作名字误杀！
        self.builtin_safe_roles = {
            "甲方", "乙方", "丙方", "丁方", "戊方", "供方", "需方", 
            "出卖人", "买受人", "委托方", "受托方", "委托人", "受托人", 
            "借款人", "贷款人", "保证人", "担保人", "担保方",
            "加工方", "承揽方", "居间方", "服务方", "顾问方", "授权方", "被授权方", "出租方", "承租方",
            "法定代表人", "法人", "法人代表", "法人或授权代表人", "授权代表", "负责人", "代理人", "联系人", "收款人", "员工", "雇员", "用人单位", "用工单位",
            "盖章", "签字", "签章", "签名"
        }
        
        # 初始化外部词库
        self.init_dictionaries()
        
        self.root.drop_target_register(DND_FILES)
        self.root.dnd_bind('<<Drop>>', self.handle_drop)
        
        # UI 元素
        self.title_label = ttk.Label(root, text="📄 Word 自动脱敏工具", font=(self.sys_font, 16, "bold"))
        self.title_label.pack(pady=5)
        
        # 叙叙专属优化：修改为精简直白的新文案提示
        self.desc_label = ttk.Label(root, text="纯本地运行，无须担心信息泄露\n仅适用于 docx 格式 | 可自定义黑名单/白名单", foreground="gray", justify="center")
        self.desc_label.pack(pady=5)
        
        # 可视化的拖拽方框
        self.drop_zone = tk.Label(
            root, text="📥\n将 Word 文档拖拽到此方框内\n（或点击此处选择文件）", 
            bg="#F0F8FF", fg="#555555", relief="solid", borderwidth=1, 
            font=(self.sys_font, 11), cursor="hand2", height=4
        )
        self.drop_zone.pack(pady=10, fill='x', ipady=15)
        self.drop_zone.bind("<Button-1>", lambda e: self.select_file())
        self.drop_zone.drop_target_register(DND_FILES)
        self.drop_zone.dnd_bind('<<Drop>>', self.handle_drop)
        
        self.file_label = ttk.Label(root, text="未选择文件", foreground="blue", wraplength=450)
        self.file_label.pack(pady=5)
        
        # 按钮容器1：开始脱敏
        style = ttk.Style()
        style.configure("Big.TButton", font=(self.sys_font, 14, "bold"))
        
        self.start_btn = ttk.Button(root, text="🚀 开始智能脱敏", command=self.start_processing_check, state=tk.DISABLED, style="Big.TButton")
        self.start_btn.pack(fill='x', pady=(15, 10), ipady=12)
        
        # 按钮容器2：黑白名单管理
        dict_frame = tk.Frame(root)
        dict_frame.pack(fill='x', pady=0)
        
        self.dict_btn_white = ttk.Button(dict_frame, text="🛡️ 管理白名单 (防误杀)", command=lambda: self.open_dictionary(self.white_dict_path))
        self.dict_btn_white.pack(side=tk.LEFT, expand=True, fill='x', padx=(0, 5))
        
        self.dict_btn_black = ttk.Button(dict_frame, text="🎯 管理黑名单 (强制杀)", command=lambda: self.open_dictionary(self.black_dict_path))
        self.dict_btn_black.pack(side=tk.RIGHT, expand=True, fill='x', padx=(5, 0))
        
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        self.status_label = ttk.Label(root, textvariable=self.status_var, font=(self.sys_font, 10))
        self.status_label.pack(pady=10)
        
        # ------------------ 授权状态底部区域 ------------------
        footer_frame = tk.Frame(root)
        footer_frame.pack(side=tk.BOTTOM, fill='x', pady=(5, 10))
        
        # 左侧显示状态
        self.license_label = tk.Label(footer_frame, font=(self.sys_font, 10))
        self.license_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # 右侧激活按钮
        self.active_btn = ttk.Button(footer_frame, text="🔑 激活软件", command=self.open_activation_window)
        self.active_btn.pack(side=tk.RIGHT, padx=(0, 10))
        
        # 底部出品标识
        self.footer_label = tk.Label(root, text="法师工具实验室 出品", font=(self.sys_font, 9), fg="#AAAAAA")
        self.footer_label.pack(side=tk.BOTTOM, pady=(0, 0))
        
        self.refresh_license_ui()

    def refresh_license_ui(self):
        """刷新授权界面展示"""
        if self.license_mgr.is_activated:
            self.license_label.config(text="✅ 软件已永久激活", fg="green")
            self.active_btn.pack_forget() # 隐藏激活按钮
        else:
            if self.license_mgr.days_left > 0:
                self.license_label.config(text=f"免费试用期，剩余 {self.license_mgr.days_left} 天", fg="#D2691E")
            else:
                self.license_label.config(text="❌ 试用期已结束，请激活", fg="red")

    def open_activation_window(self):
        """打开激活窗口"""
        act_win = tk.Toplevel(self.root)
        act_win.title("激活软件")
        act_win.geometry("380x280")
        act_win.configure(padx=20, pady=20)
        act_win.grab_set() # 模态窗口
        
        tk.Label(act_win, text="请将下方的【机器码】发送给开发者获取激活码", font=(self.sys_font, 10), fg="gray").pack(pady=(0, 10))
        
        # 机器码区域
        machine_frame = tk.Frame(act_win)
        machine_frame.pack(fill='x', pady=5)
        tk.Label(machine_frame, text="您的机器码：", font=(self.sys_font, 10, "bold")).pack(anchor='w')
        
        machine_code_var = tk.StringVar(value=self.license_mgr.get_machine_code())
        machine_entry = ttk.Entry(machine_frame, textvariable=machine_code_var, state='readonly', font=("Consolas", 11))
        machine_entry.pack(side=tk.LEFT, fill='x', expand=True, ipady=3)
        
        def copy_machine_code():
            self.root.clipboard_clear()
            self.root.clipboard_append(machine_code_var.get())
            messagebox.showinfo("复制成功", "机器码已复制，请发送给开发者！", parent=act_win)
            
        ttk.Button(machine_frame, text="复制", command=copy_machine_code, width=6).pack(side=tk.RIGHT, padx=(5,0))
        
        # 激活码区域
        tk.Label(act_win, text="请输入激活码：", font=(self.sys_font, 10, "bold")).pack(anchor='w', pady=(15, 5))
        input_code_var = tk.StringVar()
        code_entry = ttk.Entry(act_win, textvariable=input_code_var, font=("Consolas", 11))
        code_entry.pack(fill='x', ipady=3)
        
        def do_activate():
            code = input_code_var.get().strip()
            if not code:
                messagebox.showwarning("提示", "请输入激活码！", parent=act_win)
                return
            if self.license_mgr.verify_and_activate(code):
                messagebox.showinfo("激活成功", "感谢您的使用，软件已永久激活！", parent=act_win)
                self.refresh_license_ui()
                act_win.destroy()
            else:
                messagebox.showerror("激活失败", "激活码不正确，请检查是否输入有误！", parent=act_win)
                
        ttk.Button(act_win, text="🌟 立即激活", command=do_activate).pack(pady=(20, 0), fill='x', ipady=5)

    def start_processing_check(self):
        """核心卡点：开始前检查是否过期"""
        if not self.license_mgr.is_activated and self.license_mgr.days_left <= 0:
            messagebox.showwarning("试用结束", "您的 7 天免费试用期已结束。\n请点击界面右下角的【激活软件】获取正式授权码！")
            self.open_activation_window()
            return
            
        self.start_processing()

    # --------------- 下方为原有完整脱敏逻辑 ----------------

    def init_dictionaries(self):
        # 1. 创建白名单 (使用 utf-8-sig 防止记事本编码错乱)
        if not os.path.exists(self.white_dict_path):
            default_fuzzy = [
                "加工费", "运费", "材料费", "人工费", "手续费", "代理费", "服务费", "咨询费", "评估费", "保费", "差旅费", "管理费", "测试费",
                "违约", "违约金", "赔偿", "赔偿金", "诉讼", "诉讼费", "仲裁", "仲裁费", "保全", "保全费", "公证", "公证费", "鉴定", "鉴定费", "管辖", "法院", "检察院", "公安", "派出所", "法庭", "判决", "裁定", "调解", "和解", "执行", "查封", "扣押", "冻结", "抗辩", "追索", "撤销", "解除", "争议", "纠纷", "涉案", "案件", "刑事", "刑事案", "案底", "记录", "违法", "犯罪", "行政处罚", "传票", "律师", "律师费", "法务",
                "甲方", "乙方", "丙方", "丁方", "戊方", "买方", "卖方", "出让方", "受让方", "出租方", "承租方", "加工方", "承揽方", "居间方", "顾问方", "服务方", "供方", "需方", "委托方", "受托方", "委托人", "受托人",
                "合同", "协议", "备忘录", "意向书", "订单", "附件", "条款", "章节", "标的", "履行", "终止", "效力", "生效", "无效", "免责", "不可抗力", "通知", "送达", "签章", "签署", "签订", "盖章", "授权", "许可", "权利", "义务", "承诺", "保证", "声明", "陈述", "保密", "竞业限制", "所有权", "用益物权", "担保", "抵押", "质押", "留置", "定金", "保证金", "违约责任", "不可撤销", "连带责任", "印发", "出具", "发文", "发布", "签发", "批复",
                "收益", "付款", "支付", "结算", "款项", "账款", "应收", "应付", "欠款", "罚款", "金额", "比例", "费率", "发票", "开票", "税率", "增值税", "所得税", "个人所得税", "印花税", "汇率", "账户", "账号", "帐号", "帐户", "开户行", "银行", "转账", "汇款", "垫付", "首付", "预付", "尾款", "对账", "审计", "报表", "资产", "负债", "利润", "融资", "贷款", "利息", "本金", "本息", "贴现", "税号", "纳税人识别号", "统一社会信用代码", "税点", "税费", "税款", "税额", "个税", "纳税", "退税", "免税", "含税", "不含税", "代扣代缴", "汇算清缴", "应纳税",
                "公司", "企业", "集团", "机构", "董事会", "股东会", "监事会", "法定代表人", "法人", "执行董事", "总经理", "高管", "决议", "章程", "营业执照", "注册", "注销", "吊销", "破产", "清算", "并购", "重组", "股权", "股份", "增资", "减资", "分红", "收益权", "表决权", "知情权", "处分权", "优先购买权", "优先受偿权", "代位权", "债转股", "投资者", "原股东",
                "知识产权", "商标", "商标权", "专利", "专利权", "版权", "著作权", "表演者权", "邻接权", "发表权", "署名权", "修改权", "保护作品完整权", "复制权", "发行权", "出租权", "展览权", "表演权", "放映权", "广播权", "信息网络传播权", "摄制权", "改编权", "翻译权", "汇编权", "商业秘密", "集成电路布图设计权", "植物新品种权", "源码", "代码", "软件", "系统", "域名",
                "肖像权", "姓名权", "名称权", "名誉权", "荣誉权", "隐私权", "生命权", "身体权", "健康权", "物权", "债权", "期权",
                "第三方", "品牌方", "合作方", "直播", "直播平台", "公会", "MCN", "达人", "博主", "主播", "粉丝", "游客", "抖加", "投流", "推流", "刷票", "流量", "变现", "转化", "推广", "运营", "算法", "数据", "接口", "交付", "测试", "上线", "验收", "运维", "托管", "服务器", "云服务", "昵称", "实名", "演艺", "经纪", "坑位费", "佣金", "提成", "分润", "打赏", "礼物", "带货", "基数", "试生产", "小实验", "中实验", "派生",
                "税务局", "税务", "工商局", "工商", "海关", "政府", "监管", "管理局", "委员会", "仲裁委", "网信办", "工信部", "市场监督",
                "快手", "腾讯", "抖音", "小红书", "B站", "哔哩哔哩", "微博", "知乎", "斗鱼", "虎牙", "陌陌", "探探", "映客", "微视", "火山", "头条", "西瓜视频", "淘宝", "天猫", "京东", "拼多多", "优酷", "爱奇艺", "微信", "微信号", "QQ", "QQ号", "支付宝", "百度", "网易", "搜狐", "新浪",
                "吸毒", "涉毒", "涉赌", "涉黄", "酒驾", "醉驾", "台独", "港独", "疆独", "藏独", "出轨", "嫖娼", "劣迹", "绯闻", "丑闻"
            ]
            
            default_exact = [
                "LPR", "OA", "ERP", "CRM", "API", "App", "债转股",
                "台方", "平台", "平台方", "对方", "者权", "刑事案",
                ")等", "）等", "等)", "等）", ")为", "）为", ")的", "）的", ")是", "）是", ")及", "）及",
                "中华人民共和国", "中国", "香港", "澳门", "台湾", "香港特别行政区", "澳门特别行政区", 
                "特别行政区", "特区", "大陆", "中国大陆", "内地", "中国内地",
                "税", "税点", "税费"
            ]
            
            try:
                with open(self.white_dict_path, 'w', encoding='utf-8-sig') as f:
                    f.write("# =========================================================\n")
                    f.write("# 智能脱敏工具 - 白名单词库 (免死金牌)\n")
                    f.write("# 说明：此处词汇在脱敏时会被安全放过，防止误伤业务词汇。\n")
                    f.write("# =========================================================\n\n")
                    
                    f.write("[模糊匹配]\n")
                    f.write("# 说明：只要包含以下词汇，程序就会放行（例如添加“账款”，则“应收账款”也会安全）\n")
                    for word in default_fuzzy:
                        f.write(word + "\n")
                        
                    f.write("\n[精确匹配]\n")
                    f.write("# 说明：必须完全等于以下词汇，程序才会放行（适用于单字或标点碎片，防误伤）\n")
                    for word in default_exact:
                        f.write(word + "\n")
            except Exception as e:
                print(f"创建白名单文件失败: {e}")

        # 2. 创建黑名单 (使用 utf-8-sig 防止记事本编码错乱)
        if not os.path.exists(self.black_dict_path):
            try:
                with open(self.black_dict_path, 'w', encoding='utf-8-sig') as f:
                    f.write("# =========================================================\n")
                    f.write("# 智能脱敏工具 - 黑名单词库 (死亡笔记)\n")
                    f.write("# 说明：无论在文档哪里出现这里的词，都会被强制替换为【隐藏】！\n")
                    f.write("# =========================================================\n\n")
                    f.write("[强制隐藏]\n")
                    f.write("# 请在下方逐行输入你需要强制隐藏的人名、公司名或项目名：\n")
                    f.write("张三测试用名\n")
            except Exception as e:
                print(f"创建黑名单文件失败: {e}")

    def load_dictionaries(self):
        """兼容读取各种乱七八糟编码的词库，防止记事本格式暗坑"""
        self.safe_keywords = set()
        self.exact_safe_keywords = set()
        self.black_keywords = set()
        
        # 将护体金牌也注入底层的安全名单
        self.safe_keywords.update(self.builtin_safe_roles)
        
        # 【核心底层强保名单】直接注入内存！
        hardcoded_safe_words = {
            "加工", "加工费", "加工费用", "运费", "材料费", "人工费", "手续费", "代理费", "服务费", "咨询费", "评估费", "保费", "差旅费", "管理费", "研发费", "测试费",
            "违约金", "赔偿金", "保证金", "定金", "押金", "租金", "本金", "利息", "本息", "税费", "税款", "税额", "个税", "开票", "发票", "款项", "账款",
            "电脑端", "手机端", "移动端", "网页端", "客户端", "服务端", "前端", "后端", "微信端", "小程序",
            "LPR", "OA", "ERP", "CRM", "API", "App",
            # 叙叙专属：IT/系统/云服务类名词强保
            "云主机", "云平台", "云服务", "云服务器", "服务器", "主机", "虚拟机", "杀毒", "防病毒", "防火墙", "系统", "软件", "硬件", "终端", "数据库", "网络", "平台",
            "初稿", "终稿", "草案", "正本", "副本", "原件", "复印件",
            "债转股", "增资", "减资", "投资者", "投资方", "原股东", "新股东", "股东"
        }
        self.safe_keywords.update(hardcoded_safe_words)
        self.exact_safe_keywords.update(["LPR", "OA", "ERP", "CRM", "API", "App", "初稿", "终稿", "草案", "正本", "副本", "原件", "复印件", "债转股"])
        
        def read_file_lines(filepath):
            if not os.path.exists(filepath): return []
            try:
                with open(filepath, 'r', encoding='utf-8-sig') as f:
                    return f.readlines()
            except UnicodeDecodeError:
                try:
                    with open(filepath, 'r', encoding='gbk', errors='ignore') as f:
                        return f.readlines()
                except Exception:
                    return []

        white_lines = read_file_lines(self.white_dict_path)
        current_section = None
        for line in white_lines:
            line = line.strip()
            if not line or line.startswith('#'): continue
            clean_tag = line.replace(" ", "").replace("\t", "")
            if clean_tag == "[模糊匹配]":
                current_section = "fuzzy"
                continue
            if clean_tag == "[精确匹配]":
                current_section = "exact"
                continue
            if current_section == "fuzzy": self.safe_keywords.add(line)
            elif current_section == "exact": self.exact_safe_keywords.add(line)

        black_lines = read_file_lines(self.black_dict_path)
        current_section = None
        for line in black_lines:
            line = line.strip()
            if not line or line.startswith('#'): continue
            clean_tag = line.replace(" ", "").replace("\t", "")
            if clean_tag == "[强制隐藏]":
                current_section = "black"
                continue
            if current_section == "black": 
                self.black_keywords.add(line)

    def open_dictionary(self, dict_path):
        if not os.path.exists(dict_path):
            self.init_dictionaries()
            
        try:
            if sys.platform == "win32":
                os.startfile(dict_path)
            elif sys.platform == "darwin": # macOS
                subprocess.call(["open", dict_path])
            else: # linux
                subprocess.call(["xdg-open", dict_path])
        except Exception as e:
            messagebox.showerror("打开失败", f"无法自动打开词库文件，请手动打开：\n{dict_path}")

    def handle_drop(self, event):
        file_path = event.data
        if file_path.startswith('{') and file_path.endswith('}'):
            file_path = file_path[1:-1]
            
        if file_path.lower().endswith('.docx'):
            self.file_path = file_path
            self.file_label.config(text=f"已选择: {os.path.basename(file_path)}")
            self.start_btn.config(state=tk.NORMAL)
            self.drop_zone.config(bg="#E8F5E9")
            self.root.after(1000, lambda: self.drop_zone.config(bg="#F0F8FF"))
        else:
            messagebox.showwarning("格式错误", "请拖拽 Word 文档 (.docx)！")

    def select_file(self):
        filepath = filedialog.askopenfilename(
            title="选择要脱敏的 Word 文档",
            filetypes=[("Word 文档", "*.docx")]
        )
        if filepath:
            self.file_path = filepath
            self.file_label.config(text=f"已选择: {os.path.basename(filepath)}")
            self.start_btn.config(state=tk.NORMAL)

    def extract_dynamic_targets(self, doc):
        self.dynamic_black_keywords = set()
        
        target_prefixes = ["甲方", "乙方", "丙方", "丁方", "戊方", "供方", "需方", 
                           "出卖人", "买受人", "委托方", "受托方", "委托人", "受托人", 
                           "借款人", "贷款人", "保证人", "担保人", 
                           "加工方", "承揽方", "居间方", "服务方", "顾问方", "授权方", "被授权方", "出租方", "承租方"]
                           
        def add_to_memory(raw_target):
            if not raw_target: return
            
            short_name_match = re.search(r'(?:以下)?简称[为：:]?\s*[“"\'‘]?([^”"\'’）\)]+)[”"\'’]?', raw_target)
            if short_name_match:
                short_name = short_name_match.group(1).strip()
                core_short_name = re.sub(r'^[^\u4e00-\u9fa5a-zA-Z0-9]+|[^\u4e00-\u9fa5a-zA-Z0-9]+$', '', short_name)
                if short_name not in target_prefixes and len(short_name) > 1 and short_name not in ["公司", "企业", "平台", "机构"] and core_short_name not in self.builtin_safe_roles:
                    self.dynamic_black_keywords.add(short_name)
            
            def filter_parentheses(m):
                content = m.group(1)
                # 叙叙专属修复：在这里同样添加盖章签字过滤，阻止它进入废话剥离逻辑被提取
                remove_kws = ["简称", "下称", "单称", "合称", "甲方", "乙方", "丙方", "丁方", "委托", "受托", "买受", "出卖", "证", "号", "代码", "统一信用", "即", "以下", "作为", "下同", "承租", "出租", "借款", "贷款", "担保", "保证", "加工", "承揽", "居间", "服务", "顾问", "员工", "雇员", "盖章", "签字", "签章", "签名"]
                if any(kw in content for kw in remove_kws): return ""
                if re.search(r'[0-9a-zA-Z]{5,}', content): return ""
                return m.group(0) 
                
            clean_target = re.sub(r'[（\(](.*?)[）\)]', filter_parentheses, raw_target).strip()
            
            comp_match = re.search(r'([a-zA-Z0-9\u4e00-\u9fa5（）\(\)·&\-]{2,20}(?:公司|厂|局|院|中心|合伙企业|律师事务所|分行|支行))', clean_target)
            if comp_match:
                pure_name = comp_match.group(1).strip()
                if pure_name not in self.builtin_safe_roles:
                    self.dynamic_black_keywords.add(pure_name)
            else:
                clean_target = re.split(r'[,，;；\s（\(]', clean_target)[0].strip()
                clean_target = clean_target.strip(' :：,，;；.。\t\n\r\\[\\]【】()（）')
                
                # 叙叙专属修复：补充盖章签字等动作过滤
                ignore_words = {"无", "同上", "空白", "详见附件", "见附件", "未知", "签署处", "盖章处", "初稿", "终稿", "草案", "正本", "副本", "原件", "复印件", "盖章", "签字", "签章", "签名"}
                if 2 <= len(clean_target) <= 15 and clean_target not in ignore_words and clean_target not in self.builtin_safe_roles:
                    self.dynamic_black_keywords.add(clean_target)

        for para in doc.paragraphs:
            if not para.text: continue
            for match in PREFIX_REGEX.finditer(para.text):
                raw_prefix = match.group(1)
                clean_prefix = re.sub(r'[^\u4e00-\u9fa5]', '', raw_prefix)
                if any(p in clean_prefix for p in target_prefixes):
                    add_to_memory(match.group(2).strip())
                    
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                for i in range(len(cells) - 1):
                    cell_text = cells[i].text.strip()
                    if not cell_text: continue
                    clean_prefix = re.sub(r'[^\u4e00-\u9fa5]', '', cell_text)
                    if any(p == clean_prefix for p in target_prefixes):
                        add_to_memory(cells[i+1].text.strip())

    def process_text(self, text):
        if not text.strip():
            return text
            
        protected_items = []
        def save_protected(match):
            protected_items.append(match.group(0))
            return f"【PROTECTED_LAW_{len(protected_items)-1}】"
            
        def handle_book_titles(match):
            content = match.group(0)
            law_suffixes = ["法》", "法典》", "条例》", "规定》", "办法》", "解释》"]
            if any(content.endswith(suf) for suf in law_suffixes) or ("中华人民共和国" in content and "法" in content):
                return save_protected(match)
                
            hide_kws = ["协议", "合同", "意向书", "备忘录", "声明", "承诺", "函", "订单", "规则"]
            if any(kw in content for kw in hide_kws):
                return "《【隐藏】》"
            return save_protected(match)
            
        text = re.sub(r'《[^》]+》', handle_book_titles, text)
            
        all_black_words = self.black_keywords.union(self.dynamic_black_keywords)
        sorted_black_words = sorted(list(all_black_words), key=len, reverse=True)
        for black_word in sorted_black_words:
            if not black_word: continue
            core_black_word = re.sub(r'^[^\u4e00-\u9fa5a-zA-Z0-9]+|[^\u4e00-\u9fa5a-zA-Z0-9]+$', '', black_word)
            if core_black_word in self.builtin_safe_roles:
                continue
            if black_word in text:
                text = text.replace(black_word, "【隐藏】")
        
        if HIDE_ADDRESS_REGEX.search(text):
            text = HIDE_ADDRESS_REGEX.sub(r'\1【隐藏】', text)
            
        if HIDE_SHORT_REGEX.search(text):
            text = HIDE_SHORT_REGEX.sub(r'\1【隐藏】', text)

        def company_replace_strict(match):
            target = match.group(1).strip()
            if target in self.exact_safe_keywords: 
                return match.group(0)
            
            pure_suffixes = [
                "有限责任公司", "股份有限公司", "特殊普通合伙律师事务所", "特殊普通合伙制律师事务所",
                "普通合伙律师事务所", "合伙制律师事务所", "有限合伙企业", "普通合伙企业",
                "特殊普通合伙企业", "有限公司", "股份公司", "集团公司", "合伙企业", 
                "律师事务所", "分公司", "支公司", "子公司", "总公司", "代表处", "办事处", "集团", "律所"
            ]
            
            matched_suffix = ""
            prefix_part = target
            for suf in pure_suffixes:
                if target.endswith(suf):
                    matched_suffix = suf
                    prefix_part = target[:-len(suf)]
                    break
            
            boundary_words = {
                "的", "了", "在", "是", "与", "和", "或", "就", "向", "给", "对", "让", "把", "被", "及", 
                "从", "到", "自", "去", "赴", "由", "为", "等", "拟", "将", "以", "于", "按",
                "参与", "起草", "协助", "根据", "有关", "关于", "提供", "一家", "该", "本", "某", "公司", "企业",
                "发现", "发生", "存在", "属于", "具有", "包含", "包括", "产生", "进行", "实现", "完成", "导致", "造成", "作为", "成为", "视为", "要求", "确认", "约定", "规定", "证明", "保证", "承诺", "同意", "涉及", "支付", "收取", "承担", "履行", "违反", "赔偿", "通知", "解除", "终止", "签订", "签署",
                "理解", "解释", "说明", "协商", "讨论", "知道", "知悉", "获悉", "熟悉", "掌握", "获取", "取得", "保留", "保存", "保护", "维护", "适用", "修改", "变更", "补充", "完善", "执行", "配合", "支持", "审核", "审查", "评估", "判断", "认定", "看作", "采纳", "采用", "接受", "接收", "送达", "到达", "提交", "交回", "退回", "退还", "返还", "恢复", "保障",
                "甲方", "乙方", "丙方", "丁方", "戊方", "供方", "需方", 
                "委托方", "受托方", "委托人", "受托人", "买受人", "出卖人", 
                "承揽方", "加工方", "居间方", "服务方", "顾问方", "授权方", "被授权方", "出租方", "承租方",
                "双方", "单方", "共同", "均", "均为", "系"
            }
            
            cut_pos = 0
            if prefix_part:
                doc_prefix = nlp(prefix_part)
                for token in reversed(doc_prefix):
                    if token.text in boundary_words:
                        cut_pos = token.idx + len(token.text)
                        break
            
            best_prefix = prefix_part[cut_pos:].strip()
            best_target = best_prefix + matched_suffix
            
            if best_target in pure_suffixes or len(best_target) <= 4:
                return match.group(0) 
                
            return target.replace(best_target, "【隐藏】")
            
        text = COMPANY_LAWFIRM_REGEX.sub(company_replace_strict, text)
        
        text = ID_REGEX.sub("【隐藏】", text)
        text = PHONE_REGEX.sub("【隐藏】", text)
        text = TEL_REGEX.sub("【隐藏】", text)
        text = EMAIL_REGEX.sub("【隐藏】", text)
        text = PERCENTAGE_REGEX.sub("【隐藏】", text) 
        text = MONEY_REGEX_RULES.sub("【隐藏】", text)
        text = CHINESE_MONEY_REGEX.sub("【隐藏】", text)
        text = BANK_BRANCH_REGEX.sub("【隐藏】", text)
        
        doc = nlp(text)
        entities_to_replace = []
        
        for ent in doc.ents:
            clean_text = ent.text.strip()
            core_text = re.sub(r'^[^\u4e00-\u9fa5a-zA-Z0-9]+|[^\u4e00-\u9fa5a-zA-Z0-9]+$', '', clean_text)
            core_text = re.sub(r'[（\(\)）]', '', core_text)
            
            is_numeric_heading = bool(re.match(r'^[\d\.\s、\-]+$', clean_text))
            is_builtin = core_text in self.builtin_safe_roles
            
            is_safe = is_numeric_heading or \
                      is_builtin or \
                      clean_text in self.exact_safe_keywords or \
                      any(keyword in clean_text for keyword in self.safe_keywords)
            
            if "隐藏" not in clean_text and not is_safe and ent.label_ in ["PERSON", "ORG"]:
                replacement = "【隐藏】"
                if len(ent.text) > 1:
                    entities_to_replace.append((ent.text, replacement))
        
        entities_to_replace.sort(key=lambda x: len(x[0]), reverse=True)
        
        for word, replacement in entities_to_replace:
            text = text.replace(word, replacement)
            
        for i, item in enumerate(protected_items):
            text = text.replace(f"【PROTECTED_LAW_{i}】", item)
            
        text = text.replace("【【隐藏】】", "【隐藏】")
        text = text.replace("[【隐藏】]", "【隐藏】")
            
        return text

    def update_paragraph_with_tracked_changes(self, para, original_text, new_text):
        base_rpr = None
        for run in para.runs:
            rprs = run._r.xpath('./w:rPr')
            if rprs:
                base_rpr = rprs[0]
                break

        para.clear()
        p = para._p
        
        sm = difflib.SequenceMatcher(None, original_text, new_text)
        author = "智能脱敏助手"
        date_str = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')
        
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                run = para.add_run(original_text[i1:i2])
                if base_rpr is not None:
                    run._r.insert(0, copy.deepcopy(base_rpr))
            else:
                if tag in ('replace', 'delete'):
                    del_node = OxmlElement('w:del')
                    del_node.set(qn('w:id'), str(self.track_change_id))
                    del_node.set(qn('w:author'), author)
                    del_node.set(qn('w:date'), date_str)
                    self.track_change_id += 1
                    
                    del_run = OxmlElement('w:r')
                    if base_rpr is not None:
                        del_run.append(copy.deepcopy(base_rpr))
                        
                    del_text = OxmlElement('w:delText')
                    del_text.set(qn('xml:space'), 'preserve')
                    del_text.text = original_text[i1:i2]
                    
                    del_run.append(del_text)
                    del_node.append(del_run)
                    p.append(del_node)
                    
                if tag in ('replace', 'insert'):
                    ins_node = OxmlElement('w:ins')
                    ins_node.set(qn('w:id'), str(self.track_change_id))
                    ins_node.set(qn('w:author'), author)
                    ins_node.set(qn('w:date'), date_str)
                    self.track_change_id += 1
                    
                    ins_run = OxmlElement('w:r')
                    
                    rPr = OxmlElement('w:rPr')
                    if base_rpr is not None:
                        for child in base_rpr:
                            rPr.append(copy.deepcopy(child))
                            
                    highlight = OxmlElement('w:highlight')
                    highlight.set(qn('w:val'), 'yellow')
                    rPr.append(highlight)
                    ins_run.append(rPr)
                    
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')
                    t.text = new_text[j1:j2]
                    
                    ins_run.append(t)
                    ins_node.append(ins_run)
                    p.append(ins_node)

    def start_processing(self):
        self.status_var.set("正在执行脱敏规则，请稍候...")
        self.start_btn.config(state=tk.DISABLED)
        self.dict_btn_white.config(state=tk.DISABLED)
        self.dict_btn_black.config(state=tk.DISABLED)
        self.drop_zone.config(state=tk.DISABLED) 
        
        self.load_dictionaries()
        
        thread = threading.Thread(target=self.run_redaction)
        thread.start()

    def run_redaction(self):
        try:
            self.track_change_id = 1
            doc = docx.Document(self.file_path)
            
            self.extract_dynamic_targets(doc)
            
            for para in doc.paragraphs:
                if para.text:
                    original_text = para.text
                    new_text = self.process_text(original_text)
                    if new_text != original_text:
                        self.update_paragraph_with_tracked_changes(para, original_text, new_text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text:
                                original_text = para.text
                                new_text = self.process_text(original_text)
                                if new_text != original_text:
                                    self.update_paragraph_with_tracked_changes(para, original_text, new_text)

            dir_name = os.path.dirname(self.file_path)
            base_name = os.path.basename(self.file_path)
            new_file_path = os.path.join(dir_name, "已脱敏_" + base_name)
            
            doc.save(new_file_path)
            self.root.after(0, self.finish_processing, new_file_path, True)
            
        except Exception as e:
            self.root.after(0, self.finish_processing, str(e), False)

    def finish_processing(self, result, success):
        self.start_btn.config(state=tk.NORMAL)
        self.dict_btn_white.config(state=tk.NORMAL)
        self.dict_btn_black.config(state=tk.NORMAL)
        self.drop_zone.config(state=tk.NORMAL)
        
        if success:
            self.status_var.set("脱敏完成！")
            messagebox.showinfo("成功", f"文件已成功脱敏并保存至：\n{result}")
        else:
            self.status_var.set("处理失败！")
            messagebox.showerror("错误", f"处理过程中发生错误：\n{result}")

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = RedactorApp(root)
    root.mainloop()